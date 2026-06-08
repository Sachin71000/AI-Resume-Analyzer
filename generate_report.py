from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0, 51, 102)
    hs.font.name = 'Calibri'

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

# ===== TITLE PAGE =====
for _ in range(6):
    doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('Smart Resume Analyzer')
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0, 51, 102)

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run('An AI-Driven Resume-to-Job-Description Matching System\nUsing NLP, Machine Learning, and Generative AI')
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run('Case Study Report')
r3.font.size = Pt(16); r3.bold = True

doc.add_page_break()

# ===== 1. CASE STUDY TITLE =====
doc.add_heading('1. Case Study Title', level=1)
doc.add_paragraph('Smart Resume Analyzer: An AI-Driven Resume-to-Job-Description Matching System Using NLP, Machine Learning, and Generative AI')

# ===== 2. CASE STUDY OVERVIEW =====
doc.add_heading('2. Case Study Overview', level=1)
doc.add_paragraph('The Smart Resume Analyzer is a full-stack web application that automatically evaluates a candidate\'s resume against a given job description (JD). It produces a multi-dimensional compatibility score and actionable improvement suggestions using a hybrid pipeline that combines:')
doc.add_paragraph('Rule-based NLP (spaCy, regex, taxonomy matching)', style='List Bullet')
doc.add_paragraph('Classical ML (TF-IDF vectorization, cosine similarity via scikit-learn)', style='List Bullet')
doc.add_paragraph('Generative AI (Google Gemini 2.0 Flash for deep skill extraction and personalized coaching)', style='List Bullet')
doc.add_paragraph('The system addresses a real-world recruitment pain point: manually screening resumes is time-consuming and inconsistent. By automating analysis across six scoring dimensions (Semantic Similarity, Skill Match, Section Coverage, Keyword Density, Resume Quality, and ATS Compatibility), the tool provides objective, data-driven feedback to job seekers.')

doc.add_heading('Key Features', level=2)
add_table(doc, ['Feature', 'Description'], [
    ['Multi-format Parsing', 'Supports PDF and DOCX resume uploads'],
    ['Hybrid Scoring Engine', 'Weighted combination of TF-IDF, semantic, skill, section, keyword scores'],
    ['Section-aware Skill Extraction', 'Extracts skills from Skills, Projects, Experience, Education sections separately'],
    ['ATS Compatibility Scoring', '10-factor ATS audit (keywords, headings, contact, formatting, spelling, readability, bullets)'],
    ['AI-Powered Suggestions', 'Gemini 2.0 Flash generates personalized improvement advice with example rewrites'],
    ['Bullet Quality Analysis', 'Evaluates each bullet point for action verbs, metrics, and outcomes'],
    ['Experience Grading', 'Regex-based years-of-experience detection and gap analysis'],
    ['Improvement Tracking', 'Re-analyze improved resumes and compute score deltas'],
    ['Analysis History & Comparison', 'Persistent storage with side-by-side comparison'],
    ['PDF Report Export', 'Download detailed analysis as a PDF report'],
])

doc.add_heading('Technology Stack', level=2)
add_table(doc, ['Layer', 'Technology'], [
    ['Backend', 'Python, Flask, SQLAlchemy, SQLite'],
    ['NLP', 'spaCy (en_core_web_md), regex, textstat, pyspellchecker'],
    ['ML', 'scikit-learn (TF-IDF, cosine similarity)'],
    ['AI', 'Google Gemini 2.0 Flash (via google-genai SDK)'],
    ['Frontend', 'React 18, TypeScript, Vite, Framer Motion, Lucide Icons'],
    ['Data', 'Custom JSON taxonomies (500+ skills, 130+ action verbs, 60+ contextual phrases)'],
])

# ===== 3. DATASET DESCRIPTION =====
doc.add_heading('3. Dataset Description', level=1)
doc.add_paragraph('This project uses a custom-built knowledge base rather than a traditional ML dataset. The data assets are:')

doc.add_heading('3.1 Skills Taxonomy (skills_taxonomy.json)', level=2)
doc.add_paragraph('A hand-curated taxonomy of 500+ technical and soft skills organized into 8 categories:')
add_table(doc, ['Category', 'Example Skills', 'Count'], [
    ['Programming Languages', 'Python, Java, JavaScript, C++, SQL, HTML/CSS', '~40'],
    ['Frameworks & Libraries', 'React, Django, Flask, TensorFlow, PyTorch, Spring Boot', '~60'],
    ['Databases', 'MySQL, PostgreSQL, MongoDB, Redis, Elasticsearch', '~30'],
    ['Cloud & DevOps', 'AWS, Docker, Kubernetes, Terraform, CI/CD, Jenkins', '~50'],
    ['Data Science & ML', 'Machine Learning, Deep Learning, NLP, Computer Vision, LLM', '~60'],
    ['Tools', 'Git, Jira, Figma, Postman, VS Code', '~35'],
    ['Soft Skills', 'Leadership, Communication, Agile, Problem Solving', '~35'],
    ['Security', 'Cybersecurity, Penetration Testing, OWASP, Encryption', '~55'],
])

doc.add_heading('3.2 Synonyms Map (synonyms.json)', level=2)
doc.add_paragraph('Maps abbreviations and alternate names to canonical skill names (e.g., "js" → "javascript", "k8s" → "kubernetes", "ml" → "machine learning").')

doc.add_heading('3.3 Skill Contexts (skill_contexts.json)', level=2)
doc.add_paragraph('Maps contextual phrases to implied skills for inference (e.g., "phishing website detection" → ["cybersecurity", "phishing", "machine learning"]). Contains 60+ contextual patterns.')

doc.add_heading('3.4 Action Verbs (action_verbs.json)', level=2)
doc.add_paragraph('Categorized list of 130+ resume action verbs in three tiers:')
add_table(doc, ['Tier', 'Examples', 'Count'], [
    ['Strong', 'achieved, architected, engineered, spearheaded, optimized', '58'],
    ['Moderate', 'analyzed, configured, documented, researched, validated', '49'],
    ['Weak', 'helped, used, worked, tried, attended', '23'],
])

doc.add_heading('3.5 Input Data (User-Provided at Runtime)', level=2)
doc.add_paragraph('Resume: PDF or DOCX file uploaded by the user', style='List Bullet')
doc.add_paragraph('Job Description: Free-text JD pasted by the user', style='List Bullet')

# ===== 4. DATA PREPROCESSING =====
doc.add_heading('4. Data Preprocessing', level=1)
doc.add_paragraph('The preprocessing pipeline transforms raw resume and JD text into analysis-ready formats through multiple stages:')

doc.add_heading('4.1 Document Parsing (ResumeParser)', level=2)
doc.add_paragraph('Raw File (PDF/DOCX) → Plain Text + Metadata')
doc.add_paragraph('PDF: Uses pdfplumber to extract text page-by-page, preserving layout', style='List Bullet')
doc.add_paragraph('DOCX: Uses python-docx to extract paragraph text and table contents', style='List Bullet')
doc.add_paragraph('Output: { text: str, metadata: { pages, type } }', style='List Bullet')

doc.add_heading('4.2 Text Cleaning & NLP (TextPreprocessor)', level=2)
doc.add_paragraph('Raw Text → Clean Text → Tokenized + Lemmatized Text')
doc.add_paragraph('Steps performed:')
doc.add_paragraph('Entity Extraction: Regex-based extraction of emails and phone numbers', style='List Bullet')
doc.add_paragraph('Character Cleaning: Remove special characters while preserving skill-relevant symbols (/, &, -, #, +)', style='List Bullet')
doc.add_paragraph('Whitespace Normalization: Collapse multiple spaces', style='List Bullet')
doc.add_paragraph('spaCy NLP Pipeline (en_core_web_md): Lowercasing, Tokenization, Stop word removal, Punctuation removal, Lemmatization', style='List Bullet')

doc.add_heading('4.3 Section Detection (SectionDetector)', level=2)
doc.add_paragraph('Uses regex pattern matching against 10 section heading categories. Heuristic: headings are lines < 80 characters. Handles variations (e.g., "Work Experience", "Professional Experience", "Career History" all map to experience).')

doc.add_heading('4.4 Synonym Resolution (SkillExtractor)', level=2)
doc.add_paragraph('Multi-word synonyms resolved first (longest-match-first greedy strategy). Single-word synonyms resolved second. Compound splitting: "python/django" → "python django".')

# ===== 5. EDA =====
doc.add_heading('5. Exploratory Data Analysis (EDA)', level=1)
doc.add_paragraph('The system performs real-time EDA on each uploaded resume:')

doc.add_heading('5.1 Resume Structure Analysis', level=2)
add_table(doc, ['Metric', 'How It\'s Measured'], [
    ['Word Count', 'len(text.split())'],
    ['Length Status', 'too_short (<200), short (<400), good (400-1200), long (>1200)'],
    ['Section Coverage', 'Checks 3 required + 3 recommended sections'],
    ['Coverage Score', '(sections_found / 6) × 100'],
])

doc.add_heading('5.2 Content Quality Analysis', level=2)
add_table(doc, ['Metric', 'Detection Method'], [
    ['Action Verb %', 'First word of each line checked against 130-verb taxonomy'],
    ['Strong vs Weak Verbs', 'Categorized count of strong vs weak verbs'],
    ['Quantifiable Metrics', 'Regex patterns for percentages, dollars, multipliers, counts'],
    ['Contact Info', 'Regex detection of email, phone, LinkedIn, GitHub'],
    ['Bullet Structure', 'Regex detection of bullet-point markers'],
])

doc.add_heading('5.3 ATS Compatibility Audit (10-Factor)', level=2)
add_table(doc, ['Factor', 'Max Points', 'Method'], [
    ['Keyword Match', '20', 'Skill overlap percentage with JD'],
    ['Standard Headings', '10', 'Count of standard section names used'],
    ['Contact Info', '10', 'Email (4) + Phone (4) + LinkedIn (2)'],
    ['Formatting', '10', 'Tab count check, layout complexity'],
    ['Skill Placement', '5', 'Dedicated Skills section present?'],
    ['Keyword Density', '10', 'Critical missing keyword penalty'],
    ['File Parsing Quality', '10', 'Garbled character detection, text length check'],
    ['Spelling', '10', 'pyspellchecker on first 100 words > 3 chars'],
    ['Readability', '10', 'Flesch-Kincaid grade level (target: 6-16)'],
    ['Bullet Structure', '5', 'Bullet point usage in experience section'],
])

doc.add_heading('5.4 Skill Distribution Analysis', level=2)
doc.add_paragraph('Found Skills: Skills present in both resume and JD (with source tracking)', style='List Bullet')
doc.add_paragraph('Missing Skills: JD skills absent from resume, ranked by importance', style='List Bullet')
doc.add_paragraph('Resume-Only Skills: Skills in resume but not in JD', style='List Bullet')
doc.add_paragraph('Match Percentage: |found| / |jd_skills| × 100', style='List Bullet')

# ===== 6. FEATURE SELECTION =====
doc.add_heading('6. Feature Selection / Feature Engineering', level=1)

doc.add_heading('6.1 Features Used for Matching', level=2)
add_table(doc, ['Feature', 'Type', 'Weight', 'Extraction Method'], [
    ['TF-IDF Cosine Similarity', 'Numerical (0-100)', '25%', 'scikit-learn TfidfVectorizer on processed text'],
    ['Semantic Similarity', 'Numerical (0-100)', '20%', 'spaCy en_core_web_md word vectors'],
    ['Skill Match Percentage', 'Numerical (0-100)', '30%', 'Taxonomy + contextual + AI skill overlap'],
    ['Section Coverage Score', 'Numerical (0-100)', '15%', 'Required + recommended section presence'],
    ['Keyword Density Score', 'Numerical (0-100)', '10%', 'Top 15 JD keywords found in resume'],
])

doc.add_heading('6.2 Feature Engineering Techniques', level=2)
doc.add_paragraph('Multi-Source Skill Extraction: Skills extracted from 4 independent methods (taxonomy, contextual, section-aware, AI) and merged', style='List Bullet')
doc.add_paragraph('Importance Scoring: JD keyword frequency determines importance (≥3 = critical, 1-2 = important, 0 = nice-to-have)', style='List Bullet')
doc.add_paragraph('Source Priority Hierarchy: explicit > project > coursework > experience > inferred > ai', style='List Bullet')
doc.add_paragraph('Synonym Normalization: All skills canonicalized before comparison', style='List Bullet')
doc.add_paragraph('Compound Skill Splitting: "Python/Django" split into individual tokens', style='List Bullet')
doc.add_paragraph('Bullet Quality Features: Each bullet scored on action_verb + metric + outcome', style='List Bullet')
doc.add_paragraph('Experience Years Extraction: Regex patterns detect "5+ years", "3-5 years", etc.', style='List Bullet')

# ===== 7. MODEL BUILDING =====
doc.add_heading('7. Model Building', level=1)
doc.add_paragraph('The system uses three algorithmic approaches operating in a hybrid pipeline:')

doc.add_heading('7.1 Algorithm 1: TF-IDF + Cosine Similarity (Information Retrieval)', level=2)
doc.add_paragraph('Purpose: Measure lexical similarity between resume and JD.')
doc.add_paragraph('How it works:')
doc.add_paragraph('Both texts tokenized and converted to TF-IDF vectors using scikit-learn TfidfVectorizer', style='List Bullet')
doc.add_paragraph('Each term weighted by Term Frequency × Inverse Document Frequency', style='List Bullet')
doc.add_paragraph('Cosine similarity computed between the two document vectors', style='List Bullet')
doc.add_paragraph('Top 15 keywords extracted from JD vector by TF-IDF weight', style='List Bullet')
doc.add_paragraph('Keyword density = fraction of top JD keywords found in resume', style='List Bullet')
doc.add_paragraph('Strengths: Fast, interpretable, captures exact keyword matches.')
doc.add_paragraph('Limitations: Misses semantic relationships (e.g., "ML" vs "Machine Learning").')

doc.add_heading('7.2 Algorithm 2: Semantic Similarity via Word Embeddings (NLP)', level=2)
doc.add_paragraph('Purpose: Capture meaning-level similarity beyond exact keyword matches.')
doc.add_paragraph('How it works:')
doc.add_paragraph('Both texts processed through spaCy\'s en_core_web_md model (300-dim GloVe vectors)', style='List Bullet')
doc.add_paragraph('Document vectors computed as average of constituent word vectors', style='List Bullet')
doc.add_paragraph('Cosine similarity between document vectors gives semantic score', style='List Bullet')
doc.add_paragraph('Score normalized to 0-100 range', style='List Bullet')
doc.add_paragraph('Strengths: Catches synonyms and related concepts.')
doc.add_paragraph('Limitations: Averaged vectors lose word-order information.')

doc.add_heading('7.3 Algorithm 3: Generative AI — Google Gemini 2.0 Flash (LLM)', level=2)
doc.add_paragraph('Purpose: Deep contextual understanding, skill extraction with proficiency levels, and personalized suggestion generation.')
doc.add_paragraph('Three Gemini-powered operations with Pydantic schema validation:')
doc.add_paragraph('JD Parsing: Structured extraction of title, experience level, required/preferred skills, education, responsibilities', style='List Bullet')
doc.add_paragraph('Deep Skill Extraction: Finds skills that evade taxonomy matching, with proficiency levels (beginner/intermediate/expert)', style='List Bullet')
doc.add_paragraph('AI Suggestion Generation: Up to 6 personalized suggestions with concrete rewrite examples', style='List Bullet')
doc.add_paragraph('Reliability: Exponential backoff retry (up to 3 attempts). JD parsing and skill extraction run in parallel via ThreadPoolExecutor.')

# --- Image helper (used by Section 8 and 11) ---
img_dir = os.path.join(os.path.dirname(__file__), 'report_images')

def add_figure(doc, img_path, caption):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.2))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph()

# ===== 8. MODEL EVALUATION =====
doc.add_heading('8. Model Evaluation', level=1)

doc.add_heading('8.1 Scoring Metrics', level=2)
add_table(doc, ['Metric', 'Formula', 'Range', 'Purpose'], [
    ['Overall Match Score', 'Weighted sum of 5 sub-scores', '0-100', 'Primary compatibility metric'],
    ['TF-IDF Similarity', 'cosine_similarity(tfidf_resume, tfidf_jd) × 100', '0-100', 'Lexical overlap'],
    ['Semantic Similarity', 'spaCy doc.similarity() × 100', '0-100', 'Meaning-level match'],
    ['Skill Match %', '|found_skills| / |jd_skills| × 100', '0-100', 'Technical skill coverage'],
    ['ATS Score', 'Sum of 10 weighted factors', '0-100', 'ATS compatibility'],
    ['Quality Score', '100 - penalties', '0-100', 'Content quality'],
    ['Bullet Score', '(verb + metric + outcome) / 3', '0-3', 'Bullet point quality'],
])

doc.add_heading('8.2 Overall Score Formula', level=2)
p = doc.add_paragraph()
r = p.add_run('Overall = (TF-IDF × 0.25) + (Semantic × 0.20) + (Skill Match × 0.30) + (Section Coverage × 0.15) + (Keyword Density × 0.10)')
r.bold = True

doc.add_heading('8.3 ATS Scoring Breakdown (100 points)', level=2)
add_table(doc, ['Factor', 'Points', 'Evaluation Criteria'], [
    ['Keyword Match', '20', '(match_pct / 100) × 20'],
    ['Standard Headings', '10', 'min(10, standard_sections × 3.33)'],
    ['Contact Info', '10', 'Email(4) + Phone(4) + LinkedIn(2)'],
    ['Formatting', '10', 'Penalty for excessive tabs'],
    ['Skill Placement', '5', 'Dedicated Skills section = 5'],
    ['Keyword Density', '10', 'Penalty: critical_missing × 2'],
    ['File Parsing', '10', 'Penalty for garbled chars or short text'],
    ['Spelling', '10', '10 - (misspelled_count / 2)'],
    ['Readability', '10', 'FK grade 6-16 = full points'],
    ['Bullet Structure', '5', 'Bullets in experience section = 5'],
])

doc.add_heading('8.4 Evaluation Graphs', level=2)
doc.add_paragraph('The following charts visualize the model evaluation results from a sample analysis session (sachin_resume4.pdf):')

add_figure(doc, os.path.join(img_dir, '08_scoring_metrics.png'), 'Figure 8: Scoring Metrics Comparison — Bar chart showing individual metric scores with color-coded performance levels')

add_figure(doc, os.path.join(img_dir, '09_weight_distribution.png'), 'Figure 9: Feature Weight Distribution — Pie chart showing the contribution of each feature to the overall score')

add_figure(doc, os.path.join(img_dir, '12_score_composition.png'), 'Figure 10: Overall Score Composition — Weighted contribution of each feature to the final 53.2% overall score')

add_figure(doc, os.path.join(img_dir, '10_ats_breakdown.png'), 'Figure 11: ATS Compatibility Breakdown — Horizontal bar chart showing actual vs maximum points for each of the 10 ATS factors')

add_figure(doc, os.path.join(img_dir, '11_algorithm_comparison.png'), 'Figure 12: Algorithm Comparison — Multi-dimensional capability analysis of TF-IDF, spaCy Semantic, and Gemini across 6 criteria')

# ===== 9. CROSS VALIDATION =====
doc.add_heading('9. Cross Validation', level=1)

doc.add_heading('9.1 Multi-Algorithm Cross-Validation Strategy', level=2)
add_table(doc, ['Validation Approach', 'How It Works'], [
    ['Multi-Source Skill Validation', 'Skills extracted via 4 methods. Agreement across methods increases confidence.'],
    ['TF-IDF ↔ Semantic Cross-Check', 'High TF-IDF + low semantic = keyword stuffing. High semantic + low TF-IDF = paraphrasing.'],
    ['Rule-Based ↔ AI Cross-Check', 'Both engines generate suggestions independently; duplicates = high-confidence issues.'],
    ['Cache-Based Consistency', 'MD5 hash of inputs as cache key. Same inputs always produce same outputs.'],
])

doc.add_heading('9.2 Reliability Mechanisms', level=2)
add_table(doc, ['Mechanism', 'Implementation'], [
    ['API Retry with Backoff', 'Gemini calls retry 3× with exponential delays (1s, 2s)'],
    ['Graceful Degradation', 'If Gemini fails, system falls back to rule-based-only analysis'],
    ['Schema Validation', 'All Gemini responses validated against Pydantic schemas'],
    ['Parallel Execution', 'JD parsing and skill extraction run concurrently'],
    ['Result Caching', 'Identical inputs return cached results without re-computation'],
])

doc.add_heading('9.3 Improvement Tracking as Validation', level=2)
doc.add_paragraph('The /analyze/improve endpoint computes deltas when a user re-uploads an improved resume:')
doc.add_paragraph('Delta = New_Score - Previous_Score for overall, skill match, ATS, quality', style='List Bullet')
doc.add_paragraph('Skills gained and skills lost are tracked', style='List Bullet')
doc.add_paragraph('Creates a longitudinal validation loop where improvements produce positive deltas', style='List Bullet')

# ===== 10. RESULT COMPARISON TABLE =====
doc.add_heading('10. Result Comparison Table', level=1)

doc.add_heading('10.1 Algorithm Comparison', level=2)
add_table(doc, ['Criteria', 'TF-IDF + Cosine', 'Semantic (spaCy)', 'Generative AI (Gemini)'], [
    ['Approach', 'Bag-of-words, statistical', 'Word embeddings, vector avg', 'Large Language Model'],
    ['Similarity Type', 'Lexical (exact words)', 'Semantic (meaning)', 'Contextual (deep)'],
    ['Speed', 'Very Fast (~10ms)', 'Fast (~50ms)', 'Slower (~2-5s)'],
    ['Synonym Handling', 'None', 'Partial (vector proximity)', 'Full (contextual)'],
    ['Context Awareness', 'No', 'Limited', 'Full'],
    ['Interpretability', 'High (top keywords)', 'Medium', 'Medium (natural language)'],
    ['Cost', 'Free', 'Free', 'API cost per call'],
    ['Offline Capable', 'Yes', 'Yes', 'No (requires API)'],
    ['Weight in Score', '25%', '20%', 'Enriches skill data'],
])

doc.add_heading('10.2 Feature Contribution Analysis', level=2)
add_table(doc, ['Feature', 'Weight', 'Justification'], [
    ['Skill Match', '30%', 'Most direct measure of qualification fit'],
    ['TF-IDF Similarity', '25%', 'Captures document-level keyword alignment'],
    ['Semantic Similarity', '20%', 'Adds meaning-level understanding'],
    ['Section Coverage', '15%', 'Resume structure matters for ATS and recruiters'],
    ['Keyword Density', '10%', 'Ensures critical JD terms are present'],
])

# ===== 11. GRAPHS / VISUALIZATION =====
doc.add_heading('11. Graphs / Visualization', level=1)
doc.add_paragraph('The application provides rich, interactive visualizations across its React frontend. Below are screenshots captured from a live analysis session:')

add_table(doc, ['Visualization', 'Type', 'Purpose'], [
    ['Score Gauge', 'Animated circular SVG arc', 'Displays overall match score (0-100) with color coding'],
    ['Score Cards Grid', '6 metric cards (2×3 grid)', 'Shows TF-IDF, Skill Match, Section Coverage, Keyword Density, Quality, ATS scores'],
    ['Score Radar Chart', 'Radar/spider chart', 'Visual comparison of all 6 scoring dimensions'],
    ['Skill Badges', 'Color-coded badge chips', 'Green for found skills, Red for missing, with source indicators'],
    ['Improvement Roadmap', 'Prioritized card list', 'AI suggestions with expandable example rewrites'],
    ['Delta Indicators', 'Inline score change badges', '+X% in green or -X% in red for improvement tracking'],
    ['Confetti Animation', 'Canvas particle animation', 'Triggered when overall score exceeds 75%'],
])


doc.add_heading('11.1 Home Page — Resume Upload Interface', level=2)
doc.add_paragraph('The landing page features a modern dark-themed UI with a drag-and-drop resume upload zone and a text area for pasting the target job description.')
add_figure(doc, os.path.join(img_dir, '01_homepage.png'), 'Figure 1: Home Page — AI-Powered Resume Analysis upload interface')

doc.add_heading('11.2 Analysis History Dashboard', level=2)
doc.add_paragraph('The history page displays all past analyses with score bars, search functionality, and sort options. Each entry shows the filename, date, and overall match score.')
add_figure(doc, os.path.join(img_dir, '02_history.png'), 'Figure 2: Analysis History — Past analyses with score indicators and View/Delete actions')

doc.add_heading('11.3 Score Gauge & Score Cards Grid', level=2)
doc.add_paragraph('The results page prominently displays the Overall Match Score as an animated circular gauge (53% shown). Six sub-score cards show individual metrics: Semantic Similarity, Skill Match, Section Coverage, Keyword Density, Resume Quality, and ATS Score — each with color-coded progress bars.')
add_figure(doc, os.path.join(img_dir, '03_score_gauge.png'), 'Figure 3: Score Gauge (53%) with 6 Score Cards — Semantic (5), Skill Match (75), Section Coverage (50), Keyword Density (33), Quality (61), ATS (81)')

doc.add_heading('11.4 Skill Badges — Found vs Missing', level=2)
doc.add_paragraph('Skills are displayed as color-coded badges. Found skills appear in green (Java, C++, Cybersecurity), while missing skills are shown in red with importance indicators (Problem-Solving marked as missing).')
add_figure(doc, os.path.join(img_dir, '04_skills_badges.png'), 'Figure 4: Skills Found (green badges) and Missing Skills (red badges) with Improvement Insights roadmap')

doc.add_heading('11.5 Improvement Roadmap', level=2)
doc.add_paragraph('The Improvement Insights section uses a visual roadmap layout with prioritized suggestions grouped into "Immediate Fixes" (high priority), "Short-term Polish" (medium), and "Long-term Goals" (low). Each suggestion includes actionable advice.')
add_figure(doc, os.path.join(img_dir, '05_improvement_roadmap.png'), 'Figure 5: Improvement Insights — Roadmap with Immediate Fixes, Short-term Polish, and Long-term Goals')

doc.add_heading('11.6 Quality Checks & ATS Compatibility', level=2)
doc.add_paragraph('Side-by-side panels display Quality Checks (warnings about word count, action verbs, metrics) and ATS Compatibility issues (non-standard headings, missing sections, spelling errors, readability).')
add_figure(doc, os.path.join(img_dir, '06_quality_ats.png'), 'Figure 6: Quality Checks and ATS Compatibility panels with issue severity indicators')

doc.add_heading('11.7 Radar Chart & Advanced Details', level=2)
doc.add_paragraph('The Advanced Details panel reveals a radar/spider chart showing the score breakdown across 6 dimensions (Semantic Match, Skills, Sections, Keywords, Quality, ATS Fit). Below it, Resume Sections shows detected (green) and missing (red) sections, and Quality Metrics displays word count, action verb percentage, metrics count, and contact info status.')
add_figure(doc, os.path.join(img_dir, '07_radar_chart.png'), 'Figure 7: Radar Chart (Score Breakdown), Resume Sections, and Quality Metrics panel')

# ===== 12. CONCLUSION =====
doc.add_heading('12. Conclusion', level=1)

doc.add_heading('Summary of Findings', level=2)
doc.add_paragraph('The Smart Resume Analyzer demonstrates that a hybrid AI pipeline combining classical NLP, machine learning, and generative AI produces significantly more useful results than any single approach alone:')
doc.add_paragraph('TF-IDF provides fast, interpretable lexical matching but misses semantic relationships', style='List Bullet')
doc.add_paragraph('spaCy word embeddings add meaning-level understanding but lack deep contextual awareness', style='List Bullet')
doc.add_paragraph('Google Gemini provides the deepest understanding with personalized suggestions, but adds latency and API cost', style='List Bullet')

doc.add_heading('Key Achievements', level=2)
add_table(doc, ['Achievement', 'Detail'], [
    ['500+ skills in taxonomy', 'Comprehensive coverage across 8 categories'],
    ['10-factor ATS audit', 'Most thorough open-source ATS compatibility check'],
    ['3-layer skill extraction', 'Taxonomy + contextual inference + AI deep extraction'],
    ['130+ action verbs', 'Classified by strength for bullet quality analysis'],
    ['Parallel AI execution', 'JD parsing and skill extraction run concurrently'],
    ['Improvement tracking', 'Longitudinal score comparison with delta analysis'],
])

doc.add_heading('Limitations', level=2)
doc.add_paragraph('Skills taxonomy requires manual curation and may not cover niche domains', style='List Bullet')
doc.add_paragraph('Semantic similarity via averaged word vectors loses word-order information', style='List Bullet')
doc.add_paragraph('Gemini API dependency introduces latency and cost for AI features', style='List Bullet')
doc.add_paragraph('ATS scoring rules are generalized and may not perfectly match specific ATS systems', style='List Bullet')
doc.add_paragraph('Experience year detection via regex may miss non-standard date formats', style='List Bullet')

doc.add_heading('Future Enhancements', level=2)
doc.add_paragraph('Fine-tuned embedding model (e.g., sentence-transformers) for domain-specific semantic matching', style='List Bullet')
doc.add_paragraph('Date-range parsing for more accurate experience calculation', style='List Bullet')
doc.add_paragraph('Industry-specific taxonomies (healthcare, finance, legal)', style='List Bullet')
doc.add_paragraph('Multi-language support for international resume analysis', style='List Bullet')
doc.add_paragraph('Batch processing for recruiters analyzing multiple resumes against one JD', style='List Bullet')

# --- Save ---
output_path = os.path.join(os.path.dirname(__file__), 'Smart_Resume_Analyzer_Report.docx')
doc.save(output_path)
print(f"Report saved to: {output_path}")
