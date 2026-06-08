from docx import Document

# Resume
doc = Document()
doc.add_heading('John Doe - Software Engineer', 0)
doc.add_paragraph('john.doe@example.com | (123) 456-7890 | github.com/johndoe')
doc.add_heading('Summary', level=1)
doc.add_paragraph('Experienced Python developer with 5 years of experience in backend development, data analysis, and machine learning. Proficient in web frameworks like Flask and Django.')
doc.add_heading('Experience', level=1)
doc.add_heading('Senior Backend Engineer | Tech Solutions Inc.', level=2)
doc.add_paragraph('Developed RESTful APIs using Python, Flask, and PostgreSQL. Improved query performance by 40%. Delivered a machine learning model using scikit-learn.')
doc.add_heading('Skills', level=1)
doc.add_paragraph('Python, Flask, Django, SQL, PostgreSQL, Docker, AWS, Machine Learning, Data Science')
doc.save('sample_resume.docx')

# Job Description
with open('sample_jd.txt', 'w') as f:
    f.write('''
Title: Python Developer
    
We are looking for a skilled Python developer to join our backend team.
Requirements:
- 3+ years of experience with Python
- Experience building REST APIs with Flask or Django
- Strong SQL skills, ideally with PostgreSQL
- Familiarity with cloud platforms like AWS
- Bonus: Knowledge of Machine Learning algorithms
    ''')

print("Created sample_resume.docx and sample_jd.txt")
