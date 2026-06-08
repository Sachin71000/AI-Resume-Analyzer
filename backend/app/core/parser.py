import os
import pdfplumber
from docx import Document

class ResumeParser:
    @staticmethod
    def parse(file_path: str) -> dict:
        """
        Detect file type and extract text.
        Returns a dict with 'text' and 'metadata'.
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return ResumeParser._parse_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return ResumeParser._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

    @staticmethod
    def _parse_pdf(file_path: str) -> dict:
        text = ""
        pages_count = 0
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
            
        return {
            'text': text.strip(),
            'metadata': {
                'pages': pages_count,
                'type': 'pdf'
            }
        }

    @staticmethod
    def _parse_docx(file_path: str) -> dict:
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            # also extract tables if needed eventually
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
            
        return {
            'text': text.strip(),
            'metadata': {
                'type': 'docx'
            }
        }
